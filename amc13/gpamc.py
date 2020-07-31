##########################################
# PYTHON AUTOMATED MAP CHECKING ANALYSIS #
# Version: 1.01                          #
# Variant: ArcGIS Pro Geoprocessing Tool #
# Date: January 2020                     #
##########################################


# ----------------------------#
# Preliminaries and Libraries #
#-----------------------------#


# Importing the required libraries into the project
import os, arcpy, math, json, html, datetime, codecs, socket
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT



#======================#
#  MAIN Function: amc  #
#======================#

def amc(cadname, scale, scalefactor, cadpath, prjpath, tpob=None, direction=None, tolerance=2, seal = "SealKH.png", doctemplate = "LDTemplate.docx"):
    """
    Function AMC: This is the main Automated Map Checking (AMC) analysis using CAD drawings

    INPUT
        cadname: the name of the CAD drawing (without the file extension)
        scale:
        scalefactor:
        cadpath: the path to the CAD drawing (.dwg)
        prjpath: the path to the project directory
        tpob

    """

    # The output path of the project
    outpath = os.path.join(prjpath, cadname)

    # Check if the folder exist. If not, create new directory
    if not os.path.exists(outpath):
        os.makedirs(outpath)

    # Change the project's directory to the output path defined above
    os.chdir(outpath)

    # Define the project's geodatabase path
    gdbpath = os.path.join(outpath, "Reference.gdb")

    # Create a new execution report
    report = "ExecutionReport.txt"
    f = open(report, "w+", encoding = "utf8")
    f.write(f"{'EXECUTION REPORT':^80s}\n")
    f.write(f"{'County of Orange, OC Survey, Geospatial Services':^80s}\n")
    now = datetime.datetime.now().strftime("%m%d%Y %H:%M %p")
    f.write(f"{f'Python Class Execution Date and Time: {now}':^80s}\n\n")
    f.close()



    #==================== AMC Section I: Base Checks ====================#

    # Performing basic checks: imports CAD drawing and performs basic
    # layer and geometry checks.

    
    appendReport(report, f"\n{' PART 1: AMC BASE CHECKS EXECUTION ':-^80s}\n")
    stime = datetime.datetime.now().strftime("%m%d%Y %H:%M %p")
    appendReport(report, f"Script started on: {stime}\n")


    # Define new json to hold record checks
    jsonChecks = {}
    jsonChecks['LayerChecks'] = {}
    jsonChecks['BoundaryChecks'] = {}
    jsonChecks['BoundaryCorrections'] = {}
    jsonChecks['BoundaryLines'] = {}
    jsonChecks['BoundaryClosure'] = {}
    jsonChecks['GeometryCorrections'] = {}
    jsonChecks['GPSChecks'] = {}
    jsonChecks['GeodeticControlPoints'] = {}
    jsonChecks['TPOB'] = {}
    jsonChecks['Location'] = {}
    jsonChecks['MapGeometry'] = {}

    jsonControls = {}
    jsonControls['Title'] = self.cadname
    jsonControls['MapType'] = {}
    jsonControls['MapID'] = {}
    jsonControls['MapBookType'] = {}
    jsonControls['Book'] = {}
    jsonControls['Book']['No'] = '<Book No.>'
    jsonControls['Book']['Pages'] = '<Pages>'
    jsonControls['Registration'] = {}
    jsonControls['Registration']['EngCo'] = {}
    jsonControls['Registration']['EngSurveyorName'] = {}
    jsonControls['Registration']['EngSurveyorNumber'] = {}
    jsonControls['Location'] = {}
    jsonControls['Location']['Type'] = {}
    jsonControls['Location']['Name'] = {}
    jsonControls['Location']['County'] = {}
    jsonControls['Location']['State'] = 'California'
    jsonControls['GPS'] = {}
    jsonControls['Centroid'] = {}
    jsonControls['Areas'] = {}
    jsonControls['TPOB'] = {}

    jsonBoundary = {}
    jsonLegalDescription = {}

    # Tract Map, Parcel Map or Record of Survey Map
    if "TR" in cadname:
        maptype = "Tract"
        mapid = cadname.split("TR")[1]
        mapbooktype = "Miscellaneous Maps"
    elif "PM" in cadname:
        maptype = "Parcel"
        mapid = cadname.split("PM")[1]
        mapbooktype = "Parcel Maps"
    elif "RS" in cadname:
        maptype = "Record of Survey"
        mapid = cadname.split("RS")[1]
        mapbooktype = "Records of Survey"
    else:
        maptype = None
        mapid = None
        mapbooktype = None
    
    # Populate the JSON controls with the map information data for the project
    jsonControls["MapType"] = maptype
    jsonControls["MapID"] = mapid
    jsonControls["MapBookTYpe"] = mapbooktype

    # Append the report with the map information data
    appendReport(report, f"Identifying Map Characteristics from CAD drawing:\n\tMap Type: {jsonControls['MapType']}\n\tMap ID: {jsonControls['MapID']}\n\tMap Book Type: {jsonControls['MapBookType']}\n")

    # Define the project's spatial reference: NAD83 State Plane California Zone 6
    sr = arcpy.SpatialReference(102646)
    appendReport(report, "Setting Spatial Reference: NAD83 State Plane California Zone 6 (ArcGIS ID: 102646)\n")

    # Set the initial worksspace for the project's folder
    arcpy.env.workspace = outpath
    arcpy.env.OverwriteOutput = True

    # Determine if the computer is on the PFRDNET domain:
    if "PFRDNET" in socket.getfqdn():
        # Create a server geodatabase connection (for checks)
        appendReport(report, "Creating a server geodatabase connection:")
        if os.path.exists("SPOCDSQL1205.sde"):
            os.remove("SPOCDSQL1205.sde")
            appendReport(report, "\tConnection already exists: removing")
        ocserver = arcpy.CreateDatabaseConnection_management(outpath, "SPOCDSQL1205.sde", "SQL_SERVER", "10.108.9.5", "DATABASE_AUTH", "OCDataViewer", "geospatial12!", "SAVE_USERNAME")[0]
        appendReport(report, "\tConnection successfully created: SPOCDSQL1205.sde\n")
    else:
        appendReport(report, "Geodatabase connection SPOCDSQL1205.sde not found. Script outside OCPW domain.\n")


    #---------- Check 1: Check Project Geodatabase ----------#

    appendReport(report, "Project Geodatabase")
    gdbname = os.path.split(gdbpath)[1]
    appendReport(report, f"\tChecking for geodatabase: {gdbname}")

    # Check if the geodatabase exists
    if arcpy.Exists(gdbpath):
        appendReport(report, "\t...geodatabase exists.")
        # Delete the geodatabase if it exists
        arcpy.Delete_management(gdbpath)
        appendReport(report, "\t...existing geodatabase removed.")

    # Creates new geodatabase
    arcpy.CreateFileGDB_management(outpath, gdbname)
    appendReport(report, "\t...new geodatabase created.\n")


    # Change the arcpy workspace to the project's geodatabase
    arcpy.env.workspace = gdbpath
    arcpy.env.OverwriteOutput = True

    # Import the CAD drawing into the project geodatabase
    appendReport(report, "Added CAD drawing to geodatabase")
    arcpy.CADToGeodatabase_conversion(cadpath, gdbpath, cadname, "1000", sr)
    appendReport(agpmsg(1))
    appendReport(agpmsg(1))



    #---------- Check 2: Check for the presense of all the layers in CAD drawing ----------#

    # List of all the default layer types
    layerChecks = ["BASIS OF BEARING GPS TIES", "BOUNDARY", "CENTERLINES", "EASEMENTS", "LOT LINES", "NORTH ARROW MISCELLANEOUS", "RIGHT OF WAY"]

    # List of all the layer types in imported CAD drawing
    layers = []
    with arcpy.da.SearchCursor("Polyline", ["Layer"]) as cursor:
        for row in cursor:
            if row[0] not in layers:
                layers.append(row[0])
                layers.sort()

    # Perform the checks of the layers in the drawing against the defaults
    appendReport(report, "Layer Checks")
    checks = []

    for i, lyr in enumerate(layerChecks, start = 1):
        if lyr in layers:
            appendReport(report, f"\tCheck {i} of {len(layerChecks)}: {lyr} in CAD Drawing: Passed")
            jsonChecks["LayerChecks"] = "Pass"
            checks.append(True)
        elif lyr not in layers:
            appendReport(report, f"\tCheck {i} of {len(layerChecks)}: {lyr} not in CAD Drawing: Failed")
            jsonChecks["LayerChecks"] = "Fail"
            checks.append(False)

    if all(checks):
        appendReport(report, "\tAll layers passed their checks.\n")
        # Create feature classes for each of the layers in the CAD drawing in the geodatabase
        gdblayers = []
        for lyr in layerChecks:
            if lyr is not "NORTH ARROW MISCELLANEOUS":
                outfc = lyr.title().replace(" ", "")
                if arcpy.Exists(outfc):
                    arcpy.Delete_management(outfc)
                arcpy.Select_analysis("Polyline", outfc, f"""Layer = '{lyr}'""")
                gdblayers.append(outfc)
                gdblayers.sort()
    else:
        appendReport(report, "\tOne or more layers failed their checks, above. Please make sure all layers exist in the CAD drawing.\n")




    #---------- Check 3: Check for the GPS control points in the CAD drawing ----------#

    # List all of GPS points in CAD drawing and checks to make sure there are at least two of them present
    appendReport(report, "GPS Control Point Check")
    gpspoints = []
    with arcpy.da.SearchCursor("Annotation", ["RefName", "SHAPE@XY"]) as cursor:
        n = 0
        for row in cursor:
            if "GPS" in row[0]:
                n += 1
                gpspoints.append(row[0])
                jsonControls["GPS"][str(n)] = {}
                jsonControls["GPS"][str(n)]["id"] = row[0]
                jsonControls["GPS"][str(n)]["x"] = row[1][0]
                jsonControls["GPS"][str(n)]["y"] = row[1][1]

    if len(gpspoints) == 2:
        appendReport(report, "\tGPS Points Check: Passed (2 points)")
        arcpy.Select_analysis("Annotation", "GPSPoints", f"""RefName LIKE '%GPS%'""")
        appendReport(report, f"\tAdding points to geodatabase:\n {agpmsg(2)}")
        jsonChecks["GPSChecks"] = "Pass"
    elif len(gpspoints) < 2:
        appendReport(report, "\tGPS Points Check: Failed (less than 2 points)\n")
        jsonChecks["GPSChecks"] = "Fail"
    elif len(gpspoints) > 2:
        appendReport(report, "\tGPS Points Check: Failed (more than 2 points)\n")
        jsonChecks["GPSChecks"] = "Fail"




    #---------- Check 3a: Check for geodetic control geometries ----------#

    appendReport(report, "Geodetic Control Geometry Check")

    if ocserver:
        serverGC = os.path.join(ocserver, "OCSurvey.DBO.GEODETIC_HORIZONTAL")
        appendReport(report, "\tChecking Geodetic Control Server Features: OCSurvey.DBO.GEODETIC_HORIZONTAL")
    
        # Creating a new feature layer
        arcpy.MakeFeatureLayer_management(serverGC, "geodetics_lyr")

        # Updating feature class GPSPoints
        with arcpy.da.UpdateCursor("GPSPoints", ["OID@", "SHAPE@", "RefName"]) as cursor:
            for row in cursor:
                oid = row[0]
                gpsid = row[2].split("GPS NO. ")[1]

                # Searching the geodetics control layer for the GPS point geometry
                with arcpy.da.SearchCursor("geodetics_lyr", ["OID@", "SHAPE@", "GPS", "Easting2017", "Northing2017"]) as cursor1:
                    for row1 in cursor1:
                        if gpsid == row1[2]:
                            appendReport(report, f"\tGeodetic control point no. {gpsid} located in server database")
                            # Get the geometry from the server point
                            realgeometry = row1[1].WKT
                            # Transplant the geometry of the server to the geometry of the CAD layer
                            row[1] = arcpy.FromWKT(realgeometry)
                            appendReport(report, "\t\tTransplanted geometry to CAD annotation layer from server points WKT attributes")
                            # Write the coordinates to the JSON data string
                            appendReport(report, "\t\tPoint coordinates written to JSON data string")
                            jsonControls["GPS"][str(oid)]["id"] = gpsid
                            jsonControls["GPS"][str(oid)]["x"] = row1[1][0].X
                            jsonControls["GPS"][str(oid)]["y"] = row1[1][0].Y
                            cursor.updateRow(row)

        appendReport(report, "\tGeodetic Control Point Geometry Check: Passed\n")
        jsonChecks["GeodeticControlPoints"] = "Pass"

    else:
        appendReport(report, "\tChecking Geodetic Control Server Features: Failed. Script outside OCPW Domain\n")
        serverGC = None
        jsonChecks["GeodeticControlPoints"] = "Fail"




    #---------- Check 4: Check for the presence of the (True) Point of Beginning ----------#

    tpobsource = "none"
    tpobx = tpoby = 0

    # True Point of Beginning
    appendReport(report, "True Point of Beginning (TPOB) Check")

    # if TPOB coordinates are provided by user
    if tpob:
        tpobsource = "user"
        tpobx = tpob[0]
        tpoby = tpob[1]
        tpob = tpobx, tpoby
        appendReport(report, "\tTPOB provided by User: Passed\n")
        jsonChecks["TPOB"] = "Pass"
        
    # if TPOB coordinates are not provided
    elif tpob is None:
        with arcpy.da.SearchCursor("Point", ["Layer", "SHAPE@XY"]) as cursor:
            for row in cursor:
                if "TRUE POINT OF BEGINNING" in row[0]:
                    if arcpy.Exists("TPOB") == False:
                        arcpy.Select_analysis("Point", "TPOB", f"""Layer LIKE '%TRUE POINT OF BEGINNING%'""")
                        appendReport(report, "\tTPOB layer exists in CAD drawing: Passed\n")

        # If TPOB is found
        if arcpy.Exists("TPOB"):
            ntpob = int(arcpy.GetCount_management("TPOB")[0])
            if ntpob == 1:
                with arcpy.da.SearchCursor("TPOB", ["OID@", "SHAPE@"]) as cursor:
                    for row in cursor:
                        tpobsource = "cad"
                        tpobx, tpoby = row[1][0].X, row[1][0].Y
                        tpob = tpobx, tpoby
            elif ntpob > 1:
                tpobs = []
                with arcpy.da.SearchCursor("TPOB", ["OID@", "SHAPE@"]) as cursor:
                    for row in cursor:
                        tpobs.append((row[1][0].X, row[1][0].Y))

                # If all the points have the same coordinates:
                if len(tpobs) == tpobs.count(tpobs[0]):
                    msg = "WARNING: Multiple TPOB points detected in CAD Drawing. All points have the same coordinates. Using the first point in layer and ignoring the rest."
                else:
                    msg = "WARNING: Multiple TPOB points detected in CAD Drawing. These points have different coodrinates, thus selecting the correct point is untainable. This program will use the first point in the list, but you may want to re-examine the TPOB data in the original CAD drawing."

                appendReport(report, f"\t{msg}")
                tpobsource = "cad"
                tpobx, tpoby = tpobs[0][0], tpobs[0][1]
                tpob = tpobx, tpoby

        jsonChecks["TPOB"] = "Pass"
        tpobstring = "TRUE POINT OF BEGINNING"

    # If TPOB is not found
    if tpobsource == "none":
        appendReport(report, "\tTPOB is missing: Failed\n")

    # Populate the JSON Controls
    jsonControls["TPOB"]["source"] = tpobsource
    jsonControls["TPOB"]["x"] = tpobx
    jsonControls["TPOB"]["y"] = tpoby





    #---------- Check 5: Checking boundary layers ----------#

    # Checking for expanded boundary layer
    appendReport(report, "Expanded Boundary Layer Check:")
    nr = int(arcpy.GetCount_management("Boundary")[0])

    # if a single row in boundary layer
    if nr == 1:
        appendReport(report, "\tSingle boundary line detected: correcting...")
        arcpy.Rename_management("Boundary", "BoundarySingle")
        arcpy.SplitLine_management("BoundarySingle", "Boundary")
        arcpy.Delete_management("BoundarySingle")
        appendReport(report, "\tMulti-boundary lines corrected: Passed\n")
        jsonChecks["BoundaryCorrections"] = "Corrected"
        jsonChecks["BoundaryChecks"] = "Pass"
    elif nr > 1:
        appendReport(report, "\tMulti-boundary lines detected: Passed\n")
        jsonChecks["BoundaryCorrections"] = "Original"
        jsonChecks["BoundaryChecks"] = "Pass"
    else:
        appendReport(report, "\tMulti-boundary lines not detected: Failed\n")
        jsonChecks["BoundaryCorrections"] = "None"
        jsonChecks["BoundaryChecks"] = "Fail"





    #---------- Check 6: Checks for closure ----------#

    if arcpy.Exists("BoundaryArea"):
        arcpy.Delete_management("BoundaryArea")

    arcpy.FeatureToPolygon_management("Boundary", "BoundaryArea")
    appendReport(report, "Boundary Polygon/Centroid Closure Check:")

    # Adding fields
    newFields = ["CentroidX", "CentroidY", "AreaSqFeet", "AreaAcres"]
    for field in newFields:
        arcpy.AddField_management("BoundaryArea", field, "FLOAT")

    if int(arcpy.GetCount_management("BoundaryArea")[0]) > 0:
        jsonChecks["BoundaryClosure"][poid] = "Pass"
        areas = []
        # Getting the centroid coordinates for a given polygon
        appendReport(report, "\tObtaining the centroid coordinates for each boundary polygon")
        with arcpy.da.UpdateCursor("BoundaryArea", ["OID@", "SHAPE@", "CentroidX", "CentroidY", "AreaSqFeet", "AreaAcres"]) as cursor:
            for row in cursor:
                oid = row[0]
                if oid == poid: # finds the parcel OID (user provided)
                    centroidx = row[2] = row[1].centroid.X
                    centroidy = row[3] = row[1].centroid.Y
                    centroid = centroidx, centroidy # returns the centroid coordinates of that parcel
                    areaSqFeet = row[4] = row[1].getArea('GEODESIC', 'SQUAREFEET')
                    areaAcres = row[5] = row[1].getArea('GEODESIC', 'ACRES')
                    areas = (areaSqFeet, areaAcres)
                cursor.updateRow(row)

        jsonControls["Centroid"][poid] = centroid
        jsonControls["Areas"][poid] = {}
        jsonControls["Areas"][poid]["SquareFeet"] = areaSqFeet
        jsonControls["Areas"][poid]["Acres"] = areaAcres
        appendReport(report, "\tBoundary area closure: Passed\n")

    else:
        centroid = None
        jsonChecks["BoundaryClosure"][poid] = "Fail"
        appendReport(report, "\tBoundary area closure: Failed\n")





    #---------- Check 7: Checks for locations ----------#

    appendReport(report, "Map Server Location Checks")
    if ocserver:
        serverCities = os.path.join(ocserver, "OCSurvey.DBO.Boundaries\OCSurvey.DBO.CityBoundaries")
        self.appendReport(report, "\tChecking City Boundaries Server Features: OCSurvey.DBO.CityBoundaries")

        # Create temporary Cities layer from server
        arcpy.MakeFeatureLayer_management(serverCities, "cities_lyr")
        appendReport(report, "\tFiniding locations in server that intersect with CAD boundary layer (within 0.01 feet)")

        # Select cities polygons that intersect within 0.01 feet from the CAD Boundaries layer
        arcpy.SelectLayerByLocation_management("cities_lyr", "INTERSECT", "Boundary", "0.01 Feet", "NEW_SELECTION", "NOT_INVERT")

        # How many cities intersect
        citiesNo = int(arcpy.GetCount_management("cities_lyr")[0])

        # Loop the selected cities
        with arcpy.da.SearchCursor ("cities_lyr", "CITY") as cursor:
            citiesList = []
            for row in cursor:
                citiesList.append(row[0])
        if len(citiesList) == 1:
            cities = citiesList[0]
            jsonChecks["Location"] = "Pass"
            appendReport(report, f"\tLocation found: {cities}")
        elif len(citiesList) > 1:
            cities = citiesList
            jsonChecks["Location"] = "Pass"
            appendReport(report, f"\tMultiple locations found: {cities}")
        else:
            jsonChecks["Location"] = "Fail"
            appendReport(report, "\tNo locations found: Failed")

        if citiesNo == 1 and citiesList is not "UNINCORPORATED":
            cityString = f"CITY of {cities}"
            jsonControls["Location"]["Type"] = "City"
            appendReport(report, "\tLocation type: City")
        elif citiesNo ==1 and citiesList is "UNINCORPORATED":
            cityString = f"UNINCORPORATED TERRITORY"
            jsonControls["Location"]["Type"] = "Unincorporated Territory"
            appendReport(report, "\tLocation type: Unincorporated Territory")
        elif citiesNo > 1 and "UNINCORPORATED" not in citiesList:
            cityString = f"CITIES OF  {(' , ').join(tst[:-1])} AND {tst[-1]}"
            jsonControls["Location"]["type"] = "Cities"
            appendReport("\tLocation type: Cities (multiple)")
        elif citiesNo > 1 and "UNINCORPORATED" in citiesList:
            cityString = f"CITIES AND UNINCORPORATED TERRITORY OF  {(' , ').join(tst[:-1])} AND {tst[-1]}"
            jsonControls["Location"]["type"] = "Both"
            appendReport(report, "\tLocation type: Both City and Unincorporated Territory")

        jsonControls["Location"]["Name"] = cityString.title().replace("Of", "of")

        if citiesNo >= 1:
            jsonControls["Location"]["County"] = "Orange"
            countyString = "County of Orange"
            appendReport(report, f"\tCounty: {countyString}")

        appendReport(report, f"\tFull Location Identified: {cityString.title()}, {countyString}, State of California")
        if jsonChecks["Location"] == "Pass":
            appendReport(report, "\tLocation Check: Pass\n")
        elif jsonChecks["Location"] == "Fail":
            appendReport(report, "\tLocation Check: Fail\n")

    else:
        appendReport(report, "\tChecking City Boundaries from Server Features Failed: Script outside OCPW Domain.")
        citiesList, cityString = None, None




    #---------- Check 8: Check Maps (tracts, parcels or records of survey) ----------#

    if maptype == "Tract":

        appendReport(report, f"Tract Map Server Location Checks")

        if ocserver:
            serverTM = os.path.join(ocserver, "OCSurvey.DBO.TRACT_MAPS")
            appendReport(report, "\tChecking Tract Maps Server Features: OCSurvey.DBO.TRACT_MAPS")
            appendReport(report, "\tSearching Tract Map Geometry for information")

            # Create a copy of the map layer in the geodatabase:
            arcpy.MakeFeatureLayer_management(serverTM, "trackmaps_lyr")

            jsonTR = jsonControls["Title"]
            with arcpy.da.SearchCursor("trackmaps_lyr", ["OID@", "SHAPE@", "TRACTNUM", "BPNUM", "EngCo", "EngSvyName", "EngSvyNum"]) as cursor:
                for row in cursor:
                    oid = row[0]
                    if " " in row[2]:
                        serverTR = row[2].replace(" ", "")
                    else:
                        serverTR = None
                    if jsonTR == serverTR:
                        appendReport(report, f"\tServer tract lot exists in server: {jsonTR}")
                        jsonChecks["MapGeometry"] = "Pass"
                        if "/" in row[3]:
                            # if Tract then MM else if Parcel then PMB or Record of Survey then RSB
                            bb = row[3]
                            if "MM" in bb:
                                bookNo = row[3].split("/")[0].split("MM ")[1]
                            elif "PMB" in bb:
                                bookNo = row[3].split("/")[0].split("PMB ")[1]
                            elif "RSB" in bb:
                                bookNo = row[3].split("/")[0].split("PMB ")[1]
                            appendReport(report, f"\tMap book number: {bookNo}")
                            # if dash exists, split and replace dash with through, else do nothing. add inclussive at the end of string.
                            pp = row[3].split("/")[1]
                            if "-" in pp:
                                pagesNo = pp.replace("-", " through ") + " inclusive"
                            else:
                                pagesNo = row[3].split("/")[1]
                            appendReport(report, f"\tMap book pages: {pagesNo}")
                            engCo = row[4]
                            appendReport(report, f"\tEngineering Company: {engCo}")
                            engSvyName = row[5]
                            appendReport(report, f"\tSurveying Company Name: {engSvyName}")
                            engSvyNum = row[6]
                            appendReport(report, f"\tSurveying Company Number: {engSvyNum}")
                            jsonControls["Book"]["No"] = bookNo
                            jsonControls["Book"]["Pages"] = pagesNo
                            jsonControls["Registration"]["EngCo"] = engCo
                            jsonControls["Registration"]["EngSurveyorName"] = engSvyName
                            jsonControls["Registration"]["EngSurveyorNumber"] = engSvyNum
                            appendReport(report, f"\tInformation Match Found, Book No. {bookNo}, pages {pagesNo}: Passed\n")
            if jsonChecks["MapGeometry"] is not "Pass":
                appendReport(report, "\tTract map information not found in server: Failed\n")

        else:
            appendReport(report, "\tChecking Tract Maps Server Features Failed: Script outside OCPW Domain.")


    elif maptype == "Parcel":

        appendReport(report, f"Parcel Map Server Location Checks")

        if ocserver:
            serverPM = os.path.join(ocserver, "OCSurvey.DBO.PARCEL_MAPS") # Check with server for name
            appendReport(report, "\tChecking Parcel Maps Server Features: OCSurvey.DBO.PARCEL_MAPS")
            appendReport(report, "\tSearching Parcel Map Geometry for information")
            # Create a copy of the map layer in the geodatabase:
            arcpy.MakeFeatureLayer_management(serverPM, "parcelmaps_lyr")

            # Not Fully Implemented

        else:
            appendReport(report, "\tChecking Parcel Maps Server Features Failed: Script outside OCPW Domain.")


    elif maptype == "Record of Survey":

        appendReport(report, f"Tract Map Server Location Checks")

        if ocserver:
            serverRSM = os.path.join(ocserver, "OCSurvey.DBO.RECORD_OF_SURVEY") # Check with server for name
            appendReport(report, "\tChecking Record of Survey Maps Server Features: OCSurvey.DBO.RECORD_OF_SURVEY_MAPS")
            appendReport(report, "\tSearching Record of Survey Map Geometry for information")
            # Create a copy of the map layer in the geodatabase:
            arcpy.MakeFeatureLayer_management(serverRSM, "recordofsurveymaps_lyr")

            # Not Fully implemented

        else:
            appendReport(report, "\tChecking Record of Survey Maps Server Features Failed: Script outside OCPW Domain.")


    # Get the number of boundary parcels
    nParcels = jsonControls["Parcels"] = int(arcpy.GetCount_management("BoundaryArea")[0])
    appendReport(report, f"Number of parcels in boundary area: {nParcels}\n")



    #---------- Get the boundary course data ----------#

    # Create empty directionaries for the pair of lines (either direction from the point of beginning or TPOB) to be selected, and the segments of multiline coordinates and OIDs from the boundary feature class in the geodatabase
    pair = {}
    segments = {}
        
    appendReport(report, "Traverse Course Report")

    # Rounding TPOB coordinates
    rtpob = tuple([truncate(t, tolerance) for t in tpob])

    # Loop through the Boundary multilines and get OIDs and initial coordinates:
    with arcpy.da.SearchCursor("Boundary", ["OID@", "SHAPE@"]) as cursor:
        for row in cursor:
            oid = row[0] # the multiline segment OID
            start = row[1].firstPoint.X, row[1].firstPoint.Y # the initial start coordinates
            end = row[1].lastPoint.X, row[1].lastPoint.Y # the initial end coordinates

            # Update the segments dictionary to hold the segment data for each OID
            segments[oid] = {}
            segments[oid]["oid"] = oid
            segments[oid]["start"] = start
            segments[oid]["end"] = end
            segments[oid]["reversed"] = False

            # Will check later in the coe if there are results populated
            coor = None

            # Rounding start and end coordinates
            rstart = tuple([truncate(s, tolerance) for s in start])
            rend = tuple([truncate(e, tolerance) for e in end])

            # Check to see if the true point of beginning is in one of these coordinates
            if rstart == rtpob:
                coor = start, end
                reversed = False
            elif rend == rtpob:
                coor = end, start
                reversed = True

            if coor is not None:
                pair[f"{oid}"] = {}
                pair[f"{oid}"]["coor"] = coor
                pair[f"{oid}"]["reversed"] = reversed

    # Outside the arcpy row loop - choose which of the two coordinates is moving clockwise or counter-clockwise
    pts = []
    for i in pair:
        ptA = pair[i]["coor"][0]
        ptB = pair[i]["coor"][1]

        # Finds the angle degree difference from the centroid
        deg = math.degrees(math.atan2(ptA[1] - centroid[1], ptA[0] - centroid[0])) - math.degrees(math.atan2(ptB[1] - centroid[1], ptB[0] - centroid[0]))
        pts.append((i, deg))

    # Check the direction provided by the user, or if none, use clockwise direction (default)
    if direction is None or direction == "clockwise":
        appendReport(report, "\tCourse Direction: clockwise")
        # Selects the largest angle (clockwise)
        seloid = int([i[0] for i in pts if max([j[1] for j in pts]) == i[1]][0])
        selrow = pair[str(seloid)]
    elif direction == "counter-clockwise":
        appendReport(report, "\tDirection: counter-clockwise")
        # Selects the smallest angle (counter-clockwise)
        seloid = int([i[0] for i in pts if min([j[1] for j in pts]) == i[1]][0])
        selrow = pair[str(seloid)]

    # Once we select the right start line segment, we can populate the first entry of the course (with orderID = 1)
    course = {}
    course[1] = {}
    course[1]["oid"] = seloid
    course[1]["start"] = selrow["coor"][0]
    course[1]["end"] = selrow["coor"][1]
    course[1]["reversed"] = selrow["reversed"]

    # Now, given the first segment, we will run the loop for all the segments of the lines, and try to find the next start of the line (correcting at the same time the start/end coordinates of the initial feature class to make sure that start --> end follows a clockwise direction).
    while len(course) < len(segments): # runs until the course includes all the line segment
        nextLine = getnext(course, segments) # calls the getnext function above and obtains the data of the next line
        nextKey = [key for key in nextLine.keys()][0] # get the OID of the next line
        order = len(course) + 1 # update the orderID
        # Populate the next entry in the course JSON.
        course[order] = {}
        course[order]["oid"] = nextKey
        course[order]["start"] = nextLine[nextKey]["start"]
        course[order]["end"] = nextLine[nextKey]["end"]
        course[order]["reversed"] = nextLine[nextKey]["reversed"]

    # Finally, create an course order ID field (COID) in the boundary feature class and populate it with the values of the course
    arcpy.AddField_management("Boundary", "coid", "LONG", field_alias="Course Order ID")
    with arcpy.da.UpdateCursor("Boundary", ["OID@", "SHAPE@", "coid"]) as cursor:
        for row in cursor:
            oid = row[0]
            row[2] = [i for i in course if course[i]["oid"] == oid][0]
            cursor.updateRow(row)

    # Write out the course to the report
    for i in course:
        appendReport(report, f"\tCourse Order: {i}")
        appendReport(report, f"\t\tCourse OID: {self.course[i]['oid']}")
        appendReport(report, f"\t\tCourse start point: {self.course[i]['start']}")
        appendReport(report, f"\t\tCourse end point: {self.course[i]['end']}")
        appendReport(report, f"\t\tCourse reversal: {self.course[i]['reversed']}")

    # Check the size of the course
    if len(course) == len(segments):
        appendReport(report, "\tTraverse Course Complete: Passed\n")
    else:
        appendReport(report, "\tTraverse Course Incomplete: Failed\n")



    #---------- Check the boundary geometry and correct if needed ----------#

    appendReport(report, "Boundary Multiline Geometry Correction Check")

    # Update loop of the features in the geodatabase
    with arcpy.da.UpdateCursor("Boundary", ["OID@", "SHAPE@"]) as cursor:
        for row in cursor:
            oid = row[0]
            wkt = row[1].WKT
            start = row[1].firstPoint.X, row[1].firstPoint.Y
            end = row[1].lastPoint.X, row[1].lastPoint.Y
            cid = [course[i] for i in course if course[i]["oid"] == oid][0]
            if cid["start"] == start and cid["end"] == end:
                appendReport(report, f"\tOID {oid}: keeping original direction")
                rwkt = wkt
            elif cid["start"] == end and cid["end"] == start:
                appendReport(report, f"\tOID {oid}: reversing direction")
                split = wkt.split("((")[1].split("))")[0].split(", ")
                split.reverse()
                rwkt = wkt.split("((")[0] + "((" + (", ").join(split) + "))"
            geom = arcpy.FromWKT(rwkt, sr)
            row[1] = geom
            cursor.updateRow(row)

    jsonChecks["GeometryCorrections"] = "Pass"
    appendReport(report, "\tGeometry Corrections Completed: Pass\n\n")


    # Script execution update
    etime = datetime.datetime.now().strftime("%m%d%Y %H:%M %p")
    appendReport(report, f"Script completed on: {etime}\n\n")


    return



    #==================== AMC Section II: Boundary Processing ====================#

    # Processing CAD Boundaries: this session processes the boundaries of the CAD drawing
    # and performs basic checks. It also processes the boundary multiline features,
    # creates fields in the geodatabase's feature class, mathematically computes bearing,
    # distances, radial angles, etc, for annotation labels and legal descriptions.


    stime = datetime.datetime.now().strftime('%m/%d/%Y %H:%M %p')
    appendReport(report, f"\n{' PART 2: AMC BOUNDARY FEATURE PROCESSING ':-^80s}\n")
    appendReport(report, f"Script Started on: {stime}\n")

    appendReport(report, f"Processing Boundary Features for {cadname}")

    # Create fields in boundary feature class to hold types and coordinates
    boundaryFields = [["tpob", "TEXT", "", "TPOB Present"],
                        ["shapetype", "TEXT", "", "Shape Type"], 
                        ["wkt", "TEXT", "3000", "Well Known Text (WKT) Geometry"],
                        ["nwkt", "LONG", "", "Points in WKT Geometry"],
                        ["startx", "DOUBLE", "", "Startpoint X"], 
                        ["starty", "DOUBLE", "", "Startpoint Y"], 
                        ["midx", "DOUBLE", "", "Midpoint X"], 
                        ["midy", "DOUBLE", "", "Midpoint Y"], 
                        ["endx", "DOUBLE", "", "Endpoint X"], 
                        ["endy", "DOUBLE", "", "Endpoint Y"], 
                        ["midchordx", "DOUBLE", "", "Mid-chord X"],
                        ["midchordy", "DOUBLE", "", "Mid-chord Y"],
                        ["centerx", "DOUBLE", "", "Radial Center X"],
                        ["centery", "DOUBLE", "", "Radial Center Y"],
                        ["bearing", "DOUBLE", "", "Line Bearing or Chord Bearing"],
                        ["distance", "DOUBLE", "", "Line Distance or Chord Length"],
                        ["height", "DOUBLE", "", "Height of Line/Arc"],
                        ["arclength", "DOUBLE", "", "Arc Length"],
                        ["radius", "DOUBLE", "", "Arc Radius"],
                        ["midbearing", "DOUBLE", "", "Mid-chord Bearing to Center"],
                        ["delta", "DOUBLE", "", "Radial Curve Angle"],
                        ["radbearing_cs", "DOUBLE", "", "Radial Bearing: Center to Start"],
                        ["radbearing_sc", "DOUBLE", "", "Radial Bearing: Start to Center"],
                        ["radbearing_ce", "DOUBLE", "", "Radial Bearing: Center to End"],
                        ["radbearing_st", "DOUBLE", "", "Radial Tangent Angle at Start"],
                        ["radtangent", "TEXT", "", "Radial Tangent Description"],
                        ["desc_grid", "TEXT", "3000", "Legal Description (Grid)"],
                        ["desc_ground", "TEXT", "3000", "Legal Description (Ground)"],
                        ["ann_grid", "TEXT", "", "Annotation (Grid)"],
                        ["ann_ground", "TEXT", "", "Annotation (Ground)"],
                        ["annweb_grid", "TEXT", "", "Web Annotation (Grid)"],
                        ["annweb_ground", "TEXT", "", "Web Annotation (Ground)"]]

    # Add fields to Boundary feature class table
    for field in boundaryFields:
        arcpy.AddField_management("Boundary", field_name = field[0], field_type = field[1], field_length = field[2], field_alias = field[3])

    appendReport(report, f"\tAdded {len(boundaryFields)} new fields to boundary feature class")

    # Check boundary closure and populate types and coordinates
    with arcpy.da.UpdateCursor("Boundary", ["OID@", "SHAPE@"] + ["coid"] + [field[0] for field in boundaryFields]) as cursor:

        # Create empty JSON string to hold results of the loop
        jsonBoundary = {}
        # Define fields for JSON data string structure
        jsonFields = ["coid", "tpob", "shapetype", "wkt", "nwkt", "startx", "starty", "midx", "midy", "endx", "endy", "midchordx", "midchordy", "centerx", "centery", "bearing", "distance", "height", "arclength", "radius", "midbearing", "delta", "radbearing_cs", "radbearing_sc", "radbearing_ce", "radbearing_st", "radtangent", "desc_grid", "desc_ground", "ann_grid", "ann_ground", "annweb_grid", "annweb_ground"]

        # Loop through lines in feature"s multilines
        for row in cursor:
            oid = row[0]
            coid = row[2]
            jsonBoundary[oid] = {} # Indexing from OBJECTID
            
            # Create empty JSON structure for each oid
            for field in jsonFields:
                jsonBoundary[oid][field] = {}

            # Loop all fields and create an empty JSON data structure for each OID
            jsonBoundary[oid]["coid"] = coid

            #----------Current Feature----------

            # Start point coordinates
            start = startx, starty = row[1].firstPoint.X, row[1].firstPoint.Y
            row[7], row[8] = jsonBoundary[oid]["startx"], jsonBoundary[oid]["starty"] = start

            # Mid point coordinates
            mid = midx, midy = row[1].positionAlongLine(0.5, True).firstPoint.X, row[1].positionAlongLine(0.5, True).firstPoint.Y
            row[9], row[10] = jsonBoundary[oid]["midx"], jsonBoundary[oid]["midy"] = mid

            # End point coordinates
            end = endx, endy = row[1].lastPoint.X, row[1].lastPoint.Y
            row[11], row[12] = jsonBoundary[oid]["endx"], jsonBoundary[oid]["endy"] = end

            # Mid-chord coordinates
            midchord = midchordx, midchordy = (startx + endx)/2, (starty + endy)/2
            row[13], row[14] = jsonBoundary[oid]["midchordx"], jsonBoundary[oid]["midchordy"] = midchord

            # Line bearing or chord bearing
            bearing = math.degrees(math.atan2(endx - startx, endy - starty)) % 360
            row[17] = jsonBoundary[oid]["bearing"] = bearing

            # Line distance or chord length
            distance = math.hypot(endx - startx, endy - starty)
            row[18] = jsonBoundary[oid]["distance"] = distance


            # Mid-chord bearing
            midbearing = math.degrees(math.atan2(midchordx - midx, midchordy - midy)) % 360
            row[22] = jsonBoundary[oid]["midbearing"] = midbearing

            # Height of Line/Arc
            height = math.hypot(midchordx - midx, midchordy - midy)
            row[19] = jsonBoundary[oid]["height"] = height


            # Determine shape type and compute variables for lines and curves

            if height == 0: 
                # This is a line
                shapetype = "Line"
                row[4] = jsonBoundary[oid]["shapetype"] = shapetype

            elif height > 0: 
                # This is a curve
                shapetype = "Curve"
                row[4] = jsonBoundary[oid]["shapetype"] = shapetype

                # Arc radius length
                radius = (height / 2) + ((distance ** 2) / (8 * height))
                row[21] = jsonBoundary[oid]["radius"] = radius

                # Curve angle (delta)
                if height > (distance / 2): # below the diameter (more than half circle, e.g., cul-de-sac)
                    delta = (360 - math.degrees(2 * math.asin(distance / (2 * radius)))) % 360
                else: # above or at the diamerer (less or equal of half  circle)
                    delta = math.degrees( 2 * math.asin(distance / (2 * radius))) % 360
                row[23] = jsonBoundary[oid]["delta"] = delta

                # The coordinates of the center of the arc/curve
                center = centerx, centery = midx + (math.sin(math.radians(float(bearing))) * float(radius)), midy + (math.cos(math.radians(float(bearing))) * float(radius))
                row[15], row[16] = jsonBoundary[oid]["centerx"], jsonBoundary[oid]["centery"] = center

                # Tangent Check:
                radbearing_cs = math.degrees(math.atan2(startx - centerx, starty - centery)) % 360 # Radial bearing: center to start
                row[24] = jsonBoundary[oid]["radbearing_cs"] = radbearing_cs
                radbearing_sc = math.degrees(math.atan2(centerx - startx, centery - starty)) % 360 # Radial bearing: start to center
                row[25] = jsonBoundary[oid]["radbearing_sc"] = radbearing_sc
                radbearing_ce = math.degrees(math.atan2(endx - centerx, endy - centery)) % 360 # Radial bearing: center to end
                row[26] = jsonBoundary[oid]["radbearing_ce"] = radbearing_ce
                radbearing_st = (90 + radbearing_cs) % 360 # Radial Tangent Angle at start
                row[27] = jsonBoundary[oid]["radbearing_st"] = radbearing_st

                # Calculate arc length:
                arclength = (2 * math.pi * radius) * (delta / 360) # Arc Length
                row[20] = jsonBoundary[oid]["arclength"] = arclength



            # Match the TPOB with the boundary files:
            coortpob = truncate(jsonControls["TPOB"]["x"], tolerance), truncate(jsonControls["TPOB"]["y"], tolerance)
            coorstart = truncate(startx, tolerance), truncate(starty, tolerance)
            if coortpob == coorstart:
                row[3] = jsonBoundary[oid]["tpob"] = True
            else:
                row[3] = jsonBoundary[oid]["tpob"] = False

            # Well Known Text (WKT) from object"s geometry
            wkt = row[1].WKT
            row[5] = jsonBoundary[oid]["wkt"] = wkt
            
            # List and number of points in WKT
            wktlist = wkt.split("((")[1].split("))")[0].replace(" 0, ",",").replace(" 0", "").split(",")
            nwkt = len(wktlist) # Number of points in WKT geometry
            row[6] = jsonBoundary[oid]["wktcount"] = nwkt

            cursor.updateRow(row)


    appendReport(report, "\tCalculated and populated new fields in boundary feature class")


    # Repeat the same loop after writing all the previous fields and variables
    with arcpy.da.UpdateCursor("Boundary", ["OID@", "SHAPE@"] + ["coid"] + [field[0] for field in boundaryFields]) as cursor:
        for row in cursor:
            oid = row[0]
            coid = row[2]
            # Total number of features in feature class:
            nrows = int(arcpy.GetCount_management("Boundary")[0])

            # Get the last feature from the current one looping:
            if coid == 1:
                lcoid = nrows # Last coid
            elif coid > 1:
                lcoid = coid - 1

            #---------- Current Feature ----------

            # Get shapetype, start, mid, end, chordlength, midchord and height from feature attributes
            shapetype = row[4]
            start = startx, starty = row[7], row[8]
            mid = midx, midy = row[9], row[10]
            end = endx, endy = row[11], row[12]
            midchord = midchordx, midchordy = row[13], row[14]
            center = centerx, centery = row[15], row[16]
            bearing = row[17]
            distance = row[18]            
            height = row[19]
            arclength = row[20]
            radius = row[21]
            midbearing = row[22]
            delta = row[23]
            radbearing_cs = row[24]
            radbearing_sc = row[25]
            radbearing_ce = row[26]
            radbearing_st = row[27]



            # Get the preamp for the description
            if jsonBoundary[oid]["tpob"] is True:
                preamp = f"Thence from said {tpobstring}"
            else:
                preamp = " Thence"

            # Get the closing for the description
            if coid == nrows:
                closing = f" to the {tpobstring}"
            else:
                closing = ""


            # If current feature is a line:
            if shapetype == "Line":

                # Get the line description string depending on the bearing direction
                if 0 <= bearing <= 90:
                    dbearing = f"North {dd2dms(bearing)} East"
                    abearing = f"N {dd2dms(bearing)} E"
                elif 90 < bearing <= 180:
                    dbearing = f"South {dd2dms(180 - bearing)} East"
                    abearing = f"S {dd2dms(180 - bearing)} E"
                elif 180 < bearing <= 270:
                    dbearing = f"South {dd2dms(bearing - 180)} West"
                    abearing = f"S {dd2dms(bearing - 180)} W"
                elif 270 < bearing <= 360:
                    dbearing = f"North {dd2dms(360 - bearing)} West"
                    abearing = f"N {dd2dms(360 - bearing)} W"

                # Legal description (line)
                desc_grid = f"{preamp} {dbearing}, {truncate(distance, tolerance):.2f} feet;{closing}"
                desc_ground = f"{preamp} {dbearing} {truncate(distance/scalefactor, tolerance):.2f} feet;{closing}"
                row[29] = jsonBoundary[oid]["desc_grid"] = desc_grid
                row[30] = jsonBoundary[oid]["desc_ground"] = desc_ground

                # Annotation (line) for labels and web use
                ann_grid = f"{abearing}  {truncate(distance, tolerance):.2f}"
                ann_ground = f"{abearing}  {truncate(distance/scalefactor, tolerance):.2f}"
                row[31] = jsonBoundary[oid]["ann_grid"] = ann_grid
                row[32] = jsonBoundary[oid]["ann_ground"] = ann_ground
                annweb_grid = f"{abearing}\n{truncate(distance, tolerance):.2f}"
                annweb_ground = f"{abearing}\n{truncate(distance/scalefactor, tolerance):.2f}"
                row[33] = jsonBoundary[oid]["annweb_grid"] = annweb_grid
                row[34] = jsonBoundary[oid]["annweb_ground"] = annweb_ground


            # Else, if current feature is a curve:
            elif shapetype == "Curve":

                # Premp for description
                if jsonBoundary[oid]["tpob"] is False:
                    preamp = ""

                # Get the characteristics of the previous (last) feature
                with arcpy.da.SearchCursor("Boundary", ["OID@", "SHAPE@"] + ["coid"] + [field[0] for field in boundaryFields]) as lcursor:
                    for lrow in lcursor:
                        if lrow[2] == lcoid:
                            loid = lrow[0]
                            ltpob = lrow[3]
                            lshapetype = lrow[4]
                            lstart = lstartx, lstarty = lrow[7], lrow[8]
                            lmid = lmidx, lmidy = lrow[9], lrow[10]
                            lend = lendx, lendy = lrow[11], lrow[12]
                            lmidchord = lmidchordx, lmidchordy = lrow[13], lrow[14]
                            lcenter = lcenterx, lcentery = lrow[15], lrow[16]
                            lbearing = lrow[17]
                            ldistance = lrow[18]
                            lheight = lrow[19]
                            larclength = lrow[20]
                            lradius = lrow[21]
                            lmidbearing = lrow[22]
                            ldelta = lrow[23]
                            lradbearing_cs = lrow[24]
                            lradbearing_sc = lrow[25]
                            lradbearing_ce = lrow[26]
                            lradbearing_st = lrow[27]


                # Determine the last shape:
                if lshapetype == "Line":
                    # if the line bearing equals to the tangent bearing of the center-to-startpoint angle
                    if truncate(radbearing_st, tolerance) == truncate(lbearing, tolerance):
                        radtangent = "Tangent"
                    else:
                        radtangent = "Non-Tangent"

                elif lshapetype == "Curve":
                    if truncate(radbearing_cs, tolerance) == truncate(lradbearing_ce, tolerance):
                        radtangent = "Compound"
                    elif truncate(radbearing_cs, tolerance) == truncate(180 + lradbearing_ce, tolerance):
                        radtangent = "Reverse"
                    else:
                        radtangent = "Non-Tangent"

                row[28] = jsonBoundary[oid]["radtangent"] = radtangent

                # Curve Description:

                if radtangent == "Tangent":
                    desc_grid = f"{preamp} to the beginning of a curve, concave {bearing2word(midbearing)}, and having a radius of {truncate(radius, tolerance):.2f} feet; Thence {bearing2word(bearing)} along said curve {truncate(arclength, tolerance):.2f} feet through a central angle of {dd2dms(delta)};{closing}"
                    desc_ground = f"{preamp} to the beginning of a curve, concave {bearing2word(midbearing)}, and having a radius of {truncate(radius/scalefactor, tolerance):.2f} feet; Thence {bearing2word(bearing)} along said curve {truncate(arclength/scalefactor, tolerance):.2f} feet through a central angle of {dd2dms(delta)};{closing}"
                    ann_grid = f"\N{GREEK CAPITAL LETTER DELTA}={dd2dms(delta)}  R={truncate(radius, tolerance):.2f}  L={truncate(arclength, tolerance):.2f}"
                    ann_ground = f"\N{GREEK CAPITAL LETTER DELTA}={dd2dms(delta)}  R={truncate(radius/scalefactor, tolerance):.2f}  L={truncate(arclength/scalefactor, tolerance):.2f}"
                    annweb_grid = f"\N{GREEK CAPITAL LETTER DELTA}={dd2dms(delta)}\nR={truncate(radius, tolerance):.2f}\nL={truncate(arclength, tolerance):.2f}"
                    annweb_ground = f"\N{GREEK CAPITAL LETTER DELTA}={dd2dms(delta)}\nR={truncate(radius/scalefactor, tolerance):.2f}\nL={truncate(arclength/scalefactor, tolerance):.2f}"

                elif radtangent == "Compound":
                    desc_grid = f"{preamp} to the beginning of a compound curve concave {bearing2word(midbearing)} and having a radius of {truncate(radius, tolerance):.2f} feet; Thence {bearing2word(bearing)} along said curve {truncate(arclength, tolerance):.2f} feet through a central angle of {dd2dms(delta)};{closing}"
                    desc_ground = f"{preamp} to the beginning of a compound curve {bearing2word(midbearing)} and having a radius of {truncate(radius/scalefactor, tolerance):.2f} feet; Thence {bearing2word(bearing)} along said curve {truncate(arclength/scalefactor, tolerance):.2f} feet through a central angle of {dd2dms(delta)};{closing}"
                    ann_grid = f"\N{GREEK CAPITAL LETTER DELTA}={dd2dms(delta)}  R={truncate(radius, tolerance):.2f}  L={truncate(arclength, tolerance):.2f}"
                    ann_ground = f"\N{GREEK CAPITAL LETTER DELTA}={dd2dms(delta)}  R={truncate(radius/scalefactor, tolerance):.2f}  L={truncate(arclength/scalefactor, tolerance):.2f}"
                    annweb_grid = f"\N{GREEK CAPITAL LETTER DELTA}={dd2dms(delta)}\nR={truncate(radius, tolerance):.2f}\nL={truncate(arclength, tolerance):.2f}"
                    annweb_ground = f"\N{GREEK CAPITAL LETTER DELTA}={dd2dms(delta)}\nR={truncate(radius/scalefactor, tolerance):.2f}\nL={truncate(arclength/scalefactor, tolerance):.2f}"

                elif radtangent == "Reverse":
                    desc_grid = f"{preamp} to the beginning of a reverse curve concave {bearing2word(midbearing)} and having a radius of {truncate(radius, tolerance):.2f} feet; Thence {bearing2word(bearing)} along said curve {truncate(arclength, tolerance):.2f} feet through a central angle of {dd2dms(delta)};{closing}"
                    desc_ground = f"{preamp} to the beginning of a reverse curve concave {bearing2word(midbearing)} and having a radius of {truncate(radius/scalefactor, tolerance):.2f} feet; Thence {bearing2word(bearing)} along said curve {truncate(arclength/scalefactor, tolerance):.2f} feet through a central angle of {dd2dms(delta)};{closing}"
                    ann_grid = f"\N{GREEK CAPITAL LETTER DELTA}={dd2dms(delta)}  R={truncate(radius, tolerance):.2f}  L={truncate(arclength, tolerance):.2f}"
                    ann_ground = f"\N{GREEK CAPITAL LETTER DELTA}={dd2dms(delta)}  R={truncate(radius/scalefactor, tolerance):.2f}  L={truncate(arclength/scalefactor, tolerance):.2f}"
                    annweb_grid = f"\N{GREEK CAPITAL LETTER DELTA}={dd2dms(delta)}\nR={truncate(radius, tolerance):.2f}\nL={truncate(arclength, tolerance):.2f}"
                    annweb_ground = f"\N{GREEK CAPITAL LETTER DELTA}={dd2dms(delta)}\nR={truncate(radius/scalefactor, tolerance):.2f}\nL={truncate(arclength/scalefactor, tolerance):.2f}"

                elif radtangent == "Non-Tangent":
                    # Get the tangent description string depending on the bearing direction
                    if 0 <= radbearing_cs <= 90:
                        dbearing = f"North {dd2dms(radbearing_cs)} East"
                    elif 90 < radbearing_cs <= 180:
                        dbearing = f"South {dd2dms(180 - radbearing_cs)} East"
                    elif 180 < radbearing_cs <= 270:
                        dbearing = f"South {dd2dms(radbearing_cs - 180)} West"
                    elif 270 < radbearing_cs <= 360:
                        dbearing = f"North {dd2dms(360 - radbearing_cs)} West"

                    desc_grid = f"{preamp} to the beginning of a non-tangent curve, concave {bearing2word(midbearing)}, and having a radius of {truncate(radius, tolerance):.2f} feet, a radial line to said beginning of curve bears {dbearing}; Thence {bearing2word(bearing)} along said curve {truncate(arclength, tolerance):.2f} feet through a central angle of {dd2dms(delta)};{closing}"
                    desc_ground = f"{preamp} to the beginning of a non-tangent curve, concave {bearing2word(midbearing)}, and having a radius of {truncate(radius/scalefactor, tolerance):.2f} feet, a radial line to said beginning of curve bears {dbearing}; Thence {bearing2word(bearing)} along said curve {truncate(arclength/scalefactor, tolerance):.2f} feet through a central angle of {dd2dms(delta)};{closing}"
                    ann_grid = f"\N{GREEK CAPITAL LETTER DELTA}={dd2dms(delta)}  R={truncate(radius, tolerance):.2f}  L={truncate(arclength, tolerance):.2f}"
                    ann_ground = f"\N{GREEK CAPITAL LETTER DELTA}={dd2dms(delta)}  R={truncate(radius/scalefactor, tolerance):.2f}  L={truncate(arclength/scalefactor, tolerance):.2f}"
                    annweb_grid = f"\N{GREEK CAPITAL LETTER DELTA}={dd2dms(delta)}\nR={truncate(radius, tolerance):.2f}\nL={truncate(arclength, tolerance):.2f}"
                    annweb_ground = f"\N{GREEK CAPITAL LETTER DELTA}={dd2dms(delta)}\nR={truncate(radius/scalefactor, tolerance):.2f}\nL={truncate(arclength/scalefactor, tolerance):.2f}"

                # Adding the curve description to feature attributes and JSON data string
                row[29] = jsonBoundary[oid]["desc_grid"] = desc_grid
                row[30] = jsonBoundary[oid]["desc_ground"] = desc_ground
                row[31] = jsonBoundary[oid]["ann_grid"] = ann_grid
                row[32] = jsonBoundary[oid]["ann_ground"] = ann_ground
                row[33] = jsonBoundary[oid]["annweb_grid"] = annweb_grid
                row[34] = jsonBoundary[oid]["annweb_ground"] = annweb_ground

            cursor.updateRow(row)


    appendReport(report, "\tGenerated line and curve descriptions for boundary features")


    
    # Make updates and corrections
    with arcpy.da.UpdateCursor("Boundary", ["OID@", "SHAPE@"] + ["coid"] + [field[0] for field in boundaryFields]) as cursor:

        # Loop through lines in multilines
        for row in cursor:
            oid = row[0]
            coid = row[2]
            nrows = int(arcpy.GetCount_management("Boundary")[0]) # Total number of multilines
            shapetype = row[4]
            radtangent = row[28]
            desc_grid = row[29]
            desc_ground = row[30]


            # Get the previous and next multilines:
            if coid == 1:
                lcoid = nrows # Last COID
                ncoid = coid + 1 # Next COID
            elif coid > 1 and coid < nrows:
                lcoid = coid - 1
                ncoid = coid + 1
            elif coid > 1 and coid == nrows:
                lcoid = coid - 1
                ncoid = 1

            # Get the characteristics of the previous feature
            with arcpy.da.SearchCursor("Boundary", ["OID@", "SHAPE@"] + ["coid"] + [field[0] for field in boundaryFields]) as lcursor:
                for lrow in lcursor:
                    if lrow[2] == lcoid:
                        lshapetype = lrow[4]
                        lradtangent = lrow[28]
                        lradbearing_cs = lrow[24]



            # If current feature is a Line or a Non-Tangent curve and the previous is a tangent Curve:
            if shapetype == "Line" or radtangent == "Non-Tangent":
                if lshapetype == "Curve" and lradtangent == "Tangent":
                    desc_grid = desc_grid.replace("Thence", "Thence non-tangent to said curve")
                    desc_ground = desc_ground.replace("Thence", "Thence non-tangent to said curve")
                    row[29] = desc_grid
                    row[30] = desc_ground
                    jsonBoundary[oid]["desc_grid"] = desc_grid
                    jsonBoundary[oid]["desc_ground"] = desc_ground

            # If the first feature is a curve:
            if coid == 1 and shapetype == "Curve":
                #newdesc1 = f"Thence from said {tpobstring} " + desc1.split(";")[1].replace("Thence ", "")
                desc_grid = desc_grid.split(";")[1]
                desc_ground = desc_ground.split(";")[1]
                row[29] = desc_grid
                row[30] = desc_ground
                jsonBoundary[oid]["desc_grid"] = desc_grid
                jsonBoundary[oid]["desc_ground"] = desc_ground

            # if current feature is a line coming from a curve (radial)
            if shapetype == "Line" and lshapetype == "Curve":
                if bearing == lradbearing_cs or bearing == (180 + lradbearing_cs) % 360:
                    desc_grid = desc_grid.replace("Thence", "Thence radial to said curve")
                    desc_ground = desc_ground.replace("Thence", "Thence radial to said curve")
                    row[29] = desc_grid
                    row[30] = desc_ground
                    jsonBoundary[oid]["desc_grid"] = desc_grid
                    jsonBoundary[oid]["desc_ground"] = desc_ground

            cursor.updateRow(row)
    
    appendReport(report, "\tCorrected descriptions for Legal Description formatting")
    appendReport(report, "\tMultiline Descriptions added to JSON data string")

    if jsonBoundary is not None:
        appendReport(report, "\tBoundary Features Processing Complete: Passed\n")
    else:
        appendReport(report, "\tBoundary Features Processing Complete: Failed\n")

    # Write the derived annotation labels for the boundary geometry
    appendReport(report, "Annotation Labels (Grid)")
    for i in range(len(jsonBoundary)):
        jrow = [jsonBoundary[j] for j in jsonBoundary if jsonBoundary[j]["coid"] == i+1][0]
        appendReport(report, f"\tCOID {i+1} ({jrow['shapetype']}): {jrow['ann_grid'].replace('Δ', 'D')}")

    appendReport(report, "\nAnnotation Labels (Ground)")
    for i in range(len(jsonBoundary)):
        jrow = [jsonBoundary[j] for j in jsonBoundary if jsonBoundary[j]["coid"] == i+1][0]
        appendReport(report, f"\tCOID {i+1} ({jrow['shapetype']}): {jrow['ann_ground'].replace('Δ', 'D')}")

    etime = datetime.datetime.now().strftime("%m/%d/%Y %H:%M %p")
    appendReport(report, f"\nScript Completed on {etime}\n\n")





    #==================== AMC Section III: Create Legal Description ====================#

    # Create Legal Description: generates a legal description document
    # after boundary processing of the data.

    stime = datetime.datetime.now().strftime("%m/%d/%Y %H:%M %p")
    appendReport(report, f"\n{' PART 3: AMC LEGAL DESCRIPTION PROCESSING ':-^80s}\n")
    appendReport(report, f"Script Started on: {stime}\n")


    #---------- Map Document Description ----------#

    # Number of parcels in the boundary feature class
    nparcels = jsonControls["Parcels"]
    # Map type
    maptype = jsonControls["MapType"]
    if maptype == "Tract":
        kind = "Lot"
    elif maptype == "Parcel" or maptype == "Record of Survey":
        kind = "Parcel"

    if nparcels == 1:
        # If portion of lot, then use "That", else if All, use "All"
        # Need to determine if it is all or portion of a lot/parcel (TBD later)
        pre = f"That portion of {kind} 1"
    elif nparcels > 1:
        temp = []
        for i in range(nparcels):
            if i+1 < nparcels:
                temp.append(str(i+1))
            elif i+1 == nparcels:
                temp.append(f"and {i+1}")
        temp2 = ", ".join(temp)

        pre = f"These portions of {kind} {temp2}"

    # Optain the map information from the JSON Controls dataset
    loctype = jsonControls["Location"]["Type"]
    locname = jsonControls["Location"]["Name"]
    loccounty = jsonControls["Location"]["County"]
    mapid = jsonControls["MapID"]
    mapbooktype = jsonControls["MapBookType"]
    bookInfo = jsonControls["Book"]
    bookNo = bookInfo["No"]
    pagesNo = bookInfo["Pages"]

    # Generate a map description
    mapdesc = f"{pre} of {maptype} No. {mapid}, in the {locname}, County of {loccounty}, State of California, as per map filed in Book {bookNo}, pages {pagesNo} of {mapbooktype} in the Office of the County Recorder of said County, more particularly described as follows:"


    #---------- Describe Horizontal Controls ----------#

    # Create a Preamp (for Grid and Ground versions) from Horizontal Controls
    tpobx, tpoby = tpob
    gpspoints = jsonControls["GPS"]

    if len(gpspoints) == 2:
        # Calculate distances to the parcel"s centroid for each GPS point
        dist1 = math.hypot((gpspoints["1"]["x"] - centroid[0]), (gpspoints["1"]["y"] - centroid[1]))
        dist2 = math.hypot((gpspoints["2"]["x"] - centroid[0]), (gpspoints["2"]["y"] - centroid[1]))
        # Order from the most distant to the closest to the parcel points
        gpsorder = {}
        if dist1 >= dist2:
            gpsorder["HC1"] = gpspoints["1"]
            gpsorder["HC2"] = gpspoints["2"]
        elif dist1 < dist2:
            gpsorder["HC1"] = gpspoints["2"]
            gpsorder["HC2"] = gpspoints["1"]

    # GPS ID for horizontal control points
    hc1id, hc2id = gpsorder["HC1"]["id"], gpsorder["HC2"]["id"]

    # GPS coordinates for first control point (grid and ground)
    hc1x, hc1y = truncate(gpsorder["HC1"]["x"], tolerance), truncate(gpsorder["HC1"]["y"], tolerance)
    ghc1x, ghc1y = truncate(gpsorder["HC1"]["x"]/scalefactor, tolerance), truncate(gpsorder["HC1"]["y"]/scalefactor, tolerance)

    # GPS coordinates for second control point (grid and ground)
    hc2x, hc2y = truncate(gpsorder["HC2"]["x"], tolerance), truncate(gpsorder["HC2"]["y"], tolerance)
    ghc2x, ghc2y = truncate(gpsorder["HC2"]["x"]/scalefactor, tolerance), truncate(gpsorder["HC2"]["y"]/scalefactor, tolerance)
        
    # Bearing and distances (distances for both grid and ground)
    hc1bearing = math.degrees(math.atan2((hc2x - hc1x), (hc2y - hc1y))) % 360
    hc1distance = math.hypot((hc2x - hc1x), (hc2y - hc1y))
    ghc1distance = hc1distance / scalefactor
    hc2bearing = math.degrees(math.atan2((tpobx - hc2x), (tpoby - hc2y))) % 360
    hc2distance = math.hypot((tpobx - hc2x), (tpoby - hc2y))
    ghc2distance = hc2distance/ scalefactor
    annotation2 = bd2label(hc2bearing, hc2distance)
    gannotation2 = bd2label(hc2bearing, ghc2distance)

    # Get the first feature in the boundary course
    firstjson = [jsonBoundary[i] for i in jsonBoundary if jsonBoundary[i]["coid"] == 1][0]

    # if the first feature is a curve
    if firstjson["shapetype"] == "Curve":
        midbearing = firstjson["midbearing"]
        radius = firstjson["radius"]
        startx, starty = firstjson["startx"], firstjson["starty"]
        centerx, centery = firstjson["centerx"], firstjson["centery"]
        radbearing_cs = firstjson["radbearing_cs"]

        if 0 <= radbearing_cs <= 90:
            dbearing = f"North {dd2dms(radbearing_cs)} East"
        elif 90 < radbearing_cs <= 180:
            dbearing = f"South {dd2dms(180 - radbearing_cs)} East"
        elif 180 < radbearing_cs <= 270:
            dbearing = f"South {dd2dms(radbearing_cs - 180)} West"
        elif 270 < radbearing_cs <= 360:
            dbearing = f"North {dd2dms(360 - radbearing_cs)} West"

        if firstjson["radtangent"] == "Tangent":
            predesc = f", to the beginning of a curve, concave {bearing2word(midbearing)}, and having a radius of {truncate(radius, tolerance):.2f} feet, a radial bearing to said beginning of curve bears {dbearing}"
            gpredesc = f", to the beginning of a curve, concave {bearing2word(midbearing)}, and having a radius of {truncate(radius/scalefactor, tolerance):.2f} feet, a radial bearing to said beginning of curve bears {dbearing}"

        elif firstjson["radtangent"] == "Reverse":
            predesc = f", to the beginning of a reverse curve {bearing2word(midbearing)}, and having a radius of {truncate(radius, tolerance):.2f} feet, a radial bearing to said beginning of curve bears {dbearing}"
            gpredesc = f", to the beginning of a reverse curve {bearing2word(midbearing)}, and having a radius of {truncate(radius/scalefactor, tolerance):.2f} feet, a radial bearing to said beginning of curve bears {dbearing}"

        elif firstjson["radtangent"] == "Compound":
            predesc = f", to the beginning of a compound curve {bearing2word(midbearing)}, and having a radius of {truncate(radius, tolerance):.2f} feet, a radial bearing to said beginning of curve bears {dbearing}"
            gpredesc = f", to the beginning of a compound curve {bearing2word(midbearing)}, and having a radius of {truncate(radius/scalefactor, tolerance):.2f} feet, a radial bearing to said beginning of curve bears {dbearing}"

        else:
            predesc = f", to the beginning of a non-tangent curve, concave {bearing2word(midbearing)}, and having a radius of {truncate(radius, tolerance):.2f} feet, a radial bearing to said beginning of curve bears {dbearing}"
            gpredesc = f", to the beginning of a non-tangent curve, concave {bearing2word(midbearing)}, and having a radius of {truncate(radius/scalefactor, tolerance):.2f} feet, a radial bearing to said beginning of curve bears {dbearing}"

    else:
        predesc = ""
        gpredesc = ""
    
    
    preamp = f"COMMENCING at Orange County Horizontal Control Station \"{hc1id}\" having a State Plane Coordinate Value of Northing {hc1x} and Easting {hc1y}; Thence {bd2label(hc1bearing, hc1distance)} to Station \"{hc2id}\"; Thence {annotation2} to the TRUE POINT OF BEGINNING having a State Plane Coordinate Value of Northing {truncate(tpobx, tolerance):.2f} and Easting {truncate(tpoby, tolerance):.2f}{predesc}."
    gpreamp = f"COMMENCING at Orange County Horizontal Control Station \"{hc1id}\" having a State Plane Coordinate Value of Northing {ghc1x} and Easting {ghc1y}; Thence {bd2label(hc1bearing, ghc1distance)} to Station \"{hc2id}\"; Thence {gannotation2} to the TRUE POINT OF BEGINNING having a State Plane Coordinate Value of Northing {truncate(tpobx/scalefactor, tolerance):.2f} and Easting {truncate(tpoby/scalefactor, tolerance):.2f}{gpredesc}."



    #---------- Create the legal description ----------#

    ldtext = []
    gldtext = []
    for i in course:
        oid = course[i]["oid"]
        desc = jsonBoundary[oid]["desc_grid"]
        gdesc = jsonBoundary[oid]["desc_ground"]
        ldtext.append(desc)
        gldtext.append(gdesc)
    ld = "".join(ldtext)
    gld = "".join(gldtext)
    ld.replace("; to the", ", to the")
    gld.replace("; to the", ", to the")

    # Write Legal Description to Report (grid)
    appendReport(report, "\n\nLEGAL DESCRIPTION (GRID)\n")
    appendReport(report, f"\t{mapdesc}")
    appendReport(report, f"\t{preamp}")
    appendReport(report, f"\t{ld}\n")

    # Write Legal Description to Report (ground)
    appendReport(report, "\n\nLEGAL DESCRIPTION (GROUND)\n")
    appendReport(report, f"\t{mapdesc}")
    appendReport(report, f"\t{gpreamp}")
    appendReport(report, f"\t{gld}\n")

    # Compile the JSON data for the legal description
    jsonLegalDescription["Map"] = mapdesc
    jsonLegalDescription["Grid"] = {}
    jsonLegalDescription["Grid"]["Preamp"] = preamp
    jsonLegalDescription["Grid"]["Course"] = ld
    jsonLegalDescription["Ground"] = {}
    jsonLegalDescription["Ground"]["Preamp"] = gpreamp
    jsonLegalDescription["Ground"]["Course"] = gld
        
    etime = datetime.datetime.now().strftime("%m/%d/%Y %H:%M %p")
    appendReport(report, f"\nScript Completed on {etime}\n\n")




    #==================== AMC Section IV: Finalize Report ====================#

    # Finalize report and execution: compiles and exports all data and report
    # and completes the execution

    stime = datetime.datetime.now().strftime("%m/%d/%Y %H:%M %p")
    appendReport(report, f"\n{' PART 4: AMC PROCESS FINALIZATION ':-^80s}\n")
    appendReport(report, f"Script Started on: {stime}\n")

    # Compile the final JSON data
    response = {}
    response["Checks"] = jsonChecks
    response["Boundaries"] = jsonBoundary
    response["Controls"] = jsonControls
    response["LegalDescription"] = jsonLegalDescription

    os.chdir(outpath)
    with open("jsonResponse.json", "w") as jsonfile:
        json.dump(response, jsonfile)

    appendReport(report, "JSON Data String Output Written to Disk: jsonResponse.json\n")

    # Adding the legal description document to a word file
    img = os.path.join(prjpath, seal)
    template = os.path.join(prjpath, doctemplate)
    parcelid = 1


    #---------- Creating Word Document with Legal Description ----------#

    fontName = "Arial"
    fontSize = 10
    exchibitNo = "'A'"
    seal = img

    doc = Document(template)
    style = doc.styles["Normal"]
    font = style.font
    font.name = fontName
    font.size = Pt(fontSize)
    font.color.rgb = RGBColor(0, 0, 0) #Black

    # Document title
    parTitle = doc.add_heading(f"EXHIBIT {exhibitNo}")
    parTitle.style = doc.styles["Heading 1"]
    parTitle.font = parTitle.style.font
    parTitle.font.bold = True
    parTitle.font.size = Pt(fontSize + 2)
    parTitle.font.color.rgb = RGBColor(0, 0, 0) # Black
    parTitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Document heading
    parHead = doc.add_paragraph(cadname)
    parHead.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    parLDHead = doc.add_paragraph("LEGAL DESCRIPTION")
    parLDHead.style = doc.styles["Heading 2"]
    parLDHead.font = parLDHead.style.font
    parLDHead.font.bold = True
    parLDHead.font.color.rgb = RGBColor(0, 0, 0)
    parLDHead.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


    # First Paragraph (Map description and preamble)
    mapstring = jsonLegalDescription["Map"]
    doc.add_paragraph(mapstring)

    # Main Paragraph
    preampstring = jsonLegalDescription["Ground"]["Preamp"]
    p = doc.add_paragraph()
    r1= p.add_run()
    r1font = r1.font
    r1font.bold = True
    # Ties through Point of Beginning
    if "COMMENCING" in preampstring:
        r1.add_text("COMMENCING")
        part1 = preampstring.split("COMMENCING")[1]
        if "TRUE POINT OF BEGINNING" in part1:
            part2 = part1.split("TRUE POINT OF BEGINNING")
            r2 = p.add_run()
            r2.add_text(part2[0])
            r3 = p.add_run()
            r3font = r3.font
            r3font.bold = True
            r3.add_text("TRUE POINT OF BEGINNING")
            r4 = p.add_run()
            r4.add_text(part2[1])

            # Course description
            coursestring = jsonLegalDescription["Ground"]["Course"]
            part3 = coursestring.split("TRUE POINT OF BEGINNING")
            r5 = p.add_run()
            r5.add_text(part3[0].replace("; to", ", to"))
            r6 = p.add_run()
            r6font = r6.font
            r6font.bold = True
            r6.add_text("TRUE POINT OF BEGINNING.")

    # Epilogue
    areaSqFeet = int(jsonControls["Areas"][poid]["SquareFeet"])
    areaAcres = round(jsonControls["Areas"][poid]["Acres"], 3)
    doc.add_paragraph(f"Containing an area of {areaSqFeet:,} square feet, or {areaAcres:.3f} acres.")
    doc.add_paragraph()
    doc.add_paragraph(f"All values are expressed on ground values. To get the grid values multiply values by the scale factor of {scalefactor}")
    doc.add_paragraph()
    doc.add_paragraph(f"See Exhibit <Next Exhibit> attached hereto, and made a part hereof.")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("_________________________________________________________")
    doc.add_paragraph(" Kevin R. Hills, PLS 6617                                      Date")
    if seal is not None:
        doc.add_picture(seal)


    # Saving and opening the document:
    doc.save("Reference.docx")
    os.system("start Reference.docx")

    appendReport(report, "Legal Description word document created and saved to disk: Reference.docx")



    etime = datetime.datetime.now().strftime("%m/%d/%Y %H:%M %p")
    appendReport(report, f"\nScript Completed on {etime}\n\n")


    appendReport(report, f"\n{'END OF EXECUTION REPORT':^80s}\n")


    return response




#=======================#
#  SECONDARY FUNCTIONS  #
#=======================#



#---------- AMC Function: Append Report ----------#

def appendReport(report, string):
    """AMC Function: Append Execution Report"""
    # Open the file for appending
    fa = open(report, 'a+')
    # Append to the end of the file
    fa.write(f'{string}\n')
    # Close the file after appending
    fa.close()       
    print(string)
    return



#---------- AMC Function: Arcpy Message ----------#

def agpmsg(ntabs=1):
    """AMC Function: Arcpy Message"""
    tabs = "\t"*ntabs
    msg = tabs + arcpy.GetMessages().replace("\n", f"\n{tabs}") + "\n\n"
    return msg



#---------- AMC Function: Truncating Values ----------#

def truncate(v, n):
    """AMC Function: Truncating values
    Truncates coordinates at the n-th decimal places, for the value v (double)
    """
    return math.floor(v * 10 ** n) / 10 ** n





#---------- AMC Function: Get next feature in course ----------#

def getnext(course, segments):
    """AMC Function: Get next segment in boundary course
    Gets the next course coordinate based on the initial line (course[1]), and the line segment coordinates from ArcGIS Boundary Feature class. Returns a JSON string indexed by the order ID (the order to which the lines are added to the course), and for each item, the Boundary feature class OBJECTID, its true start and end coordinates (reversed from the feature class line direction if needed - always clockwise).
    """
    # This is the existing (initial) OID
    existing = [course[i]["oid"] for i in course.keys()]
    # Get the endpoint of the existing feature segment
    lastend = course[len(course)]["end"]
    result = {}
    # Search all segments to find the one whose startpoint or endpoint is the same with the last endpoint.
    for key in segments:
        if key not in existing:
            if lastend == segments[key]["start"]:
                result[key] = {}
                result[key]["start"] = segments[key]["start"]
                result[key]["end"] = segments[key]["end"]
                result[key]["reversed"] = False
            elif lastend == segments[key]["end"]:
                result[key] = {}
                result[key]["start"] = segments[key]["end"]
                result[key]["end"] = segments[key]["start"]
                result[key]["reversed"] = True
    if len(result) == 1:
        return result





#---------- AMC Function: Decimal Degrees to Degrees-Minutes-Seconds ----------#

def dd2dms(dd):
    """AMC Function: Decimal Degrees to Degrees-Minutes-Seconds.
    Returns formatted coordinates of the Degrees:Minutes:Seconds format. This secondary function restructures decimal degree coordinates into degree/minutes/seconds coordinates. It is called from the main module function.
    """
    decimaldegree = dd
    minutes = decimaldegree%1.0*60
    seconds = minutes%1.0*60
    dms = u"{0:02}\xb0{1:02}\'{2:02}'".format(int(math.floor(decimaldegree)), int(math.floor(minutes)), int(seconds))
    arcpy.AddMessage(dms)
    return dms




#---------- AMC Function: Bearing to Word ----------#

def bearing2word(bearing):
    """AMC Function: Bearing to Word
    Returns the corresponding direction word based on radial bearing values
    """
    if 0 < bearing <= 22.5:
        bWord = "northerly"
    elif 22.5 < bearing <= 67.5:
        bWord = "northeasterly"
    elif 67.5 < bearing <= 112.5:
        bWord = "easterly"
    elif 112.5 < bearing <= 157.5:
        bWord = "southeasterly"
    elif 157.5 < bearing <= 202.5:
        bWord = "southerly"
    elif 202.5 < bearing <= 247.5:
        bWord = "southwesterly"
    elif 247.5 < bearing <= 292.5:
        bWord = "westerly"
    elif 292.5 < bearing <= 337.5:
        bWord = "northwesterly"
    else:
        bWord = "northerly"

    return bWord



#---------- AMC Function: Format Labels for Bearing and Distance ----------#

def bd2label(bearing, distance):
    """AMC Function: Format Labels for Bearing and Distance
    Generates a formatted bearing and distance string from coordinates
    """
    if 0 <= bearing <= 90:
        label = f"North {dd2dms(truncate(bearing, tolerance))} East, {truncate(distance, tolerance)} feet"
    elif 90 < bearing <= 180:
        label = f"South {dd2dms(truncate(180 - bearing, tolerance))} East, {truncate(distance, tolerance)} feet"
    elif 180 < bearing <= 270:
        label = f"South {dd2dms(truncate(bearing - 180, tolerance))} West, {truncate(distance, tolerance)} feet"
    elif 270 < bearing <= 360:
        label = f"North {dd2dms(truncate(360 - bearing, tolerance))}, West, {truncate(distance, tolerance)} feet"

    return label
