##############################################################
# PYTHON AUTOMATED LEGAL DESCRIPTION DOCUMENT                #
# Version: 1.2                                               #
# Variant: ArcGIS Pro Geoprocessing Tool                     #
# Date: January 2020                                         #
##############################################################



#============================================================#
# PRELIMINARIES AND LIBRARIES                                #
#============================================================#


# Importing the required libraries into the project
import os, json, html, datetime, codecs, socket
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT



def json2ald(jsonpath, prjpath, template=None, seal=None, scale="ground", fontName="Arial", fontSize=10, exhibitNo="A", poid=1):
    """
    Function: Automated Legal Description (ALD)
        Returns a Word document containing a formatted legal description in either ground or grid coordinates.
    """
    
    # Import the JSON string
    with open(jsonpath) as file:
        jsonString = json.load(file)

    # The output path of the results (same as the directory with the JSON file)
    outpath = os.path.split(jsonpath)[0]

    # Change working directory
    os.chdir(prjpath)

    # Template check
    if template is None:
        template = os.path.join(prjpath, "LDTemplate.docx")

    # Seak check
    if seal is None:
        seal = os.path.join(prjpath, "SealKH.png")

    # Get the scale factor from the JSON string
    scalefactor = float(jsonString["Controls"]["ScaleFactor"])


    # Initializing document settings
    doc = Document(template)
    style = doc.styles["Normal"]
    font = style.font
    font.name = fontName
    font.size = Pt(fontSize)
    font.color.rgb = RGBColor(0, 0, 0) # Black color
        
    # Document title
    parTitle = doc.add_heading(f"EXHIBIT {exhibitNo}")
    parTitle.style = doc.styles["Heading 1"]
    parTitle.font = parTitle.style.font
    parTitle.font.bold = True
    parTitle.font.size = Pt(fontSize + 2)
    parTitle.font.color.rgb = RGBColor(0, 0, 0) # Black
    parTitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Document heading
    parHead = doc.add_paragraph(jsonString["Controls"]["Title"])
    parHead.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    parLDHead = doc.add_paragraph("LEGAL DESCRIPTION")
    parLDHead.style = doc.styles["Heading 2"]
    parLDHead.font = parLDHead.style.font
    parLDHead.font.bold = True
    parLDHead.font.color.rgb = RGBColor(0, 0, 0)
    parLDHead.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


    # First Paragraph (Map description and preamble)
    mapstring = jsonString["LegalDescription"]["Map"]
    doc.add_paragraph(mapstring)

    # Main Paragraph
    preampstring = jsonString["LegalDescription"][scale.capitalize()]["Preamp"]
    p = doc.add_paragraph()
    r1= p.add_run()
    r1font = r1.font
    r1font.bold = True

    # Ties through Point of Beginning
    if "COMMENCING" in preampstring:
        r1.add_text("COMMENCING")
        part1 = preampstring.split("COMMENCING")[1]
        if "TRUE POINT OF BEGINNING" in part1:
            part2 = part1.split("TRUE POINT OF BEGINNING")
            r2 = p.add_run()
            r2.add_text(part2[0])
            r3 = p.add_run()
            r3font = r3.font
            r3font.bold = True
            r3.add_text("TRUE POINT OF BEGINNING")
            r4 = p.add_run()
            r4.add_text(part2[1])

            # Course description
            coursestring = jsonString["LegalDescription"][scale.capitalize()]["Course"]
            part3 = coursestring.split("TRUE POINT OF BEGINNING")
            r5 = p.add_run()
            r5.add_text(part3[0].replace("; to", ", to"))
            r6 = p.add_run()
            r6font = r6.font
            r6font.bold = True
            r6.add_text("TRUE POINT OF BEGINNING.")

    # Epilogue
    areaSqFeet = int(jsonString["Controls"]["Areas"][str(poid)]["SquareFeet"])
    areaAcres = round(jsonString["Controls"]["Areas"][str(poid)]["Acres"], 3)
    doc.add_paragraph(f"Containing an area of {areaSqFeet:,} square feet, or {areaAcres:.3f} acres.")
    doc.add_paragraph()
    if scale == "ground":
        scalestring = f"All values are expressed on ground values. To get the grid values, multiply values by the scale factor of {scalefactor}"
    elif scale == "grid":
        scalestring = f"All values are expressed on grid values. To tet the ground values, divide values by the scale factor of {scalefactor}"
    doc.add_paragraph(scalestring)
    doc.add_paragraph()
    doc.add_paragraph(f"See Exhibit <Next Exhibit> attached hereto, and made a part hereof.")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("_________________________________________________________")
    doc.add_paragraph(" Kevin R. Hills, PLS 6617                                      Date")
    print(f"\tAdding Seal\n")
    if seal is not None:
        doc.add_picture(seal)

    # Saving and opening the document:
    os.chdir(outpath)
    doc.save("Reference.docx")

    docout = os.path.join(outpath, "Reference.docx")

    return docout


